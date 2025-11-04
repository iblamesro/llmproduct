"""
Document processing utilities for RegIntelAI
"""
import os
from typing import List, Dict, BinaryIO
from pypdf import PdfReader
from pathlib import Path
from config import CHUNK_SIZE, CHUNK_OVERLAP

# Classe simple de text splitter pour remplacer langchain
class RecursiveCharacterTextSplitter:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> list:
        """Split text into chunks"""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap
        
        return chunks


def extract_text_from_pdf(pdf_file) -> str:
    """
    Extract text from uploaded PDF file
    
    Args:
        pdf_file: Streamlit uploaded file object or file path
        
    Returns:
        Extracted text as string
    """
    try:
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += f"\n\n--- Page {page_num + 1} ---\n\n{page_text}"
        return text
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_txt(txt_file) -> str:
    """
    Extract text from TXT file
    
    Args:
        txt_file: Streamlit uploaded file object or file path
        
    Returns:
        Extracted text as string
    """
    try:
        if isinstance(txt_file, str):
            with open(txt_file, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return txt_file.read().decode('utf-8')
    except Exception as e:
        raise Exception(f"Error extracting text from TXT: {str(e)}")


def extract_text_from_docx(docx_file) -> str:
    """
    Extract text from DOCX file
    
    Args:
        docx_file: Streamlit uploaded file object or file path
        
    Returns:
        Extracted text as string
    """
    try:
        from docx import Document
        if isinstance(docx_file, str):
            doc = Document(docx_file)
        else:
            doc = Document(docx_file)
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except ImportError:
        raise Exception("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


def extract_text_from_file(file, filename: str = None) -> str:
    """
    Extract text from any supported file format
    
    Args:
        file: File object or path
        filename: Name of the file (to determine type)
        
    Returns:
        Extracted text as string
    """
    if filename is None:
        if isinstance(file, str):
            filename = file
        else:
            filename = getattr(file, 'name', '')
    
    ext = Path(filename).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file)
    elif ext == '.txt' or ext == '.md':
        return extract_text_from_txt(file)
    elif ext == '.docx':
        return extract_text_from_docx(file)
    else:
        raise Exception(f"Unsupported file format: {ext}")


def load_documents_from_folder(folder_path: str) -> List[Dict[str, str]]:
    """
    Load all documents from a folder
    
    Args:
        folder_path: Path to folder containing documents
        
    Returns:
        List of documents with text and metadata
    """
    documents = []
    folder = Path(folder_path)
    
    if not folder.exists():
        raise Exception(f"Folder not found: {folder_path}")
    
    supported_extensions = ['.pdf', '.txt', '.md', '.docx']
    
    for file_path in folder.glob('**/*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                text = extract_text_from_file(str(file_path))
                documents.append({
                    'text': text,
                    'filename': file_path.name,
                    'path': str(file_path)
                })
            except Exception as e:
                print(f"Error loading {file_path.name}: {str(e)}")
    
    return documents


def chunk_documents(text: str, filename: str) -> List[Dict[str, str]]:
    """
    Split document text into chunks with metadata
    
    Args:
        text: Document text to chunk
        filename: Name of the source file
        
    Returns:
        List of chunks with metadata
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )
    
    chunks = text_splitter.split_text(text)
    
    # Add metadata to each chunk
    chunked_docs = []
    for idx, chunk in enumerate(chunks):
        chunked_docs.append({
            "text": chunk,
            "metadata": {
                "source": filename,
                "chunk_id": idx,
                "total_chunks": len(chunks)
            }
        })
    
    return chunked_docs


def format_citations(chunks: List[Dict]) -> str:
    """
    Format retrieved chunks as citations
    
    Args:
        chunks: List of retrieved chunks with metadata
        
    Returns:
        Formatted citation string
    """
    citations = []
    for idx, chunk in enumerate(chunks):
        source = chunk.get("metadata", {}).get("source", "Unknown")
        chunk_id = chunk.get("metadata", {}).get("chunk_id", "?")
        citations.append(f"[{idx + 1}] {source} (Chunk {chunk_id + 1})")
    
    return "\n".join(citations)
